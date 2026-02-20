"""DOCX generator using python-docx.

Sprint I: Reader UI & Accessibility

Maps TipTap JSON nodes to python-docx elements using recursive traversal,
following the same pattern as tiptap_to_markdown.py.
"""

import io
import re
from typing import Any


class DocxGenerator:
    """Generates DOCX documents from TipTap JSON content."""

    @staticmethod
    def generate(
        tiptap_content: dict[str, Any],
        title: str,
        toc: list[dict[str, Any]],
        metadata: dict[str, Any] | None = None,
    ) -> tuple[io.BytesIO, str]:
        """Generate a DOCX from TipTap JSON.

        Args:
            tiptap_content: TipTap JSON document
            title: Document title
            toc: Table of contents entries
            metadata: Optional metadata

        Returns:
            Tuple of (BytesIO buffer, suggested filename)
        """
        from docx import Document
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()

        # Title
        title_para = doc.add_heading(title, level=0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Metadata
        meta = metadata or {}
        if any(meta.get(k) for k in ("version", "status", "document_number", "author")):
            meta_para = doc.add_paragraph()
            meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            parts = []
            if meta.get("document_number"):
                parts.append(meta["document_number"])
            if meta.get("version"):
                parts.append(f"Version {meta['version']}")
            if meta.get("status"):
                parts.append(f"Status: {meta['status'].title()}")
            if meta.get("author"):
                parts.append(f"Author: {meta['author']}")
            run = meta_para.add_run(" | ".join(parts))
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

        doc.add_page_break()

        # Process content
        content = tiptap_content.get("content", []) if tiptap_content else []
        for node in content:
            DocxGenerator._render_node(doc, node)

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)

        slug = re.sub(r"[^a-z0-9]+", "-", title.lower()).strip("-")
        filename = f"{slug}.docx"

        return buf, filename

    @staticmethod
    def _render_node(doc: Any, node: dict[str, Any]) -> None:
        """Render a TipTap node to DOCX elements."""
        from docx.shared import Pt

        node_type = node.get("type", "")
        attrs = node.get("attrs", {})
        content = node.get("content", [])

        if node_type == "heading":
            level = min(max(attrs.get("level", 1), 1), 6)
            # python-docx heading levels are 0-9 but 0 is Title
            heading = doc.add_heading(level=min(level, 9))
            DocxGenerator._add_inline_content(heading, content)

        elif node_type == "paragraph":
            para = doc.add_paragraph()
            DocxGenerator._add_inline_content(para, content)

        elif node_type == "bulletList":
            for item in content:
                DocxGenerator._render_list_item(doc, item, style="List Bullet")

        elif node_type == "orderedList":
            for item in content:
                DocxGenerator._render_list_item(doc, item, style="List Number")

        elif node_type == "taskList":
            for item in content:
                checked = item.get("attrs", {}).get("checked", False)
                checkbox = "[x]" if checked else "[ ]"
                para = doc.add_paragraph(style="List Bullet")
                run = para.add_run(f"{checkbox} ")
                item_content = item.get("content", [])
                for child in item_content:
                    if child.get("type") == "paragraph":
                        DocxGenerator._add_inline_content(para, child.get("content", []))

        elif node_type == "codeBlock":
            text_parts = []
            for child in content:
                if child.get("type") == "text":
                    text_parts.append(child.get("text", ""))
            code_text = "".join(text_parts)
            para = doc.add_paragraph()
            run = para.add_run(code_text)
            run.font.name = "Courier New"
            run.font.size = Pt(9)
            para.style = doc.styles["No Spacing"]

        elif node_type == "blockquote":
            for child in content:
                if child.get("type") == "paragraph":
                    para = doc.add_paragraph()
                    para.paragraph_format.left_indent = Pt(36)
                    DocxGenerator._add_inline_content(para, child.get("content", []))

        elif node_type == "horizontalRule":
            para = doc.add_paragraph()
            para.add_run("_" * 50)

        elif node_type == "table":
            DocxGenerator._render_table(doc, content)

        elif node_type == "image":
            para = doc.add_paragraph()
            alt = attrs.get("alt", "")
            src = attrs.get("src", "")
            run = para.add_run(f"[Image: {alt or src}]")
            run.italic = True

        else:
            # Unknown node type - try rendering children
            for child in content:
                DocxGenerator._render_node(doc, child)

    @staticmethod
    def _render_list_item(doc: Any, item: dict[str, Any], style: str) -> None:
        """Render a list item."""
        content = item.get("content", [])
        for child in content:
            if child.get("type") == "paragraph":
                para = doc.add_paragraph(style=style)
                DocxGenerator._add_inline_content(para, child.get("content", []))
            elif child.get("type") in ("bulletList", "orderedList"):
                # Nested lists
                for nested_item in child.get("content", []):
                    DocxGenerator._render_list_item(doc, nested_item, style=style)

    @staticmethod
    def _add_inline_content(paragraph: Any, content: list[dict[str, Any]]) -> None:
        """Add inline content (text with marks) to a paragraph."""
        from docx.shared import Pt

        for node in content:
            if node.get("type") == "text":
                text = node.get("text", "")
                marks = node.get("marks", [])
                run = paragraph.add_run(text)

                for mark in marks:
                    mark_type = mark.get("type", "")
                    if mark_type == "bold":
                        run.bold = True
                    elif mark_type == "italic":
                        run.italic = True
                    elif mark_type == "code":
                        run.font.name = "Courier New"
                        run.font.size = Pt(9)
                    elif mark_type == "strike":
                        run.font.strike = True
                    elif mark_type == "underline":
                        run.underline = True
            elif node.get("type") == "hardBreak":
                paragraph.add_run("\n")

    @staticmethod
    def _render_table(doc: Any, rows: list[dict[str, Any]]) -> None:
        """Render a table."""
        if not rows:
            return

        # Determine column count from first row
        first_row = rows[0].get("content", [])
        col_count = len(first_row)
        if col_count == 0:
            return

        table = doc.add_table(rows=len(rows), cols=col_count)
        table.style = "Table Grid"

        for row_idx, row in enumerate(rows):
            cells = row.get("content", [])
            for col_idx, cell in enumerate(cells):
                if col_idx < col_count:
                    cell_content = cell.get("content", [])
                    cell_text_parts = []
                    for para in cell_content:
                        if para.get("type") == "paragraph":
                            inline = para.get("content", [])
                            for node in inline:
                                if node.get("type") == "text":
                                    cell_text_parts.append(node.get("text", ""))
                    table.cell(row_idx, col_idx).text = "".join(cell_text_parts)
